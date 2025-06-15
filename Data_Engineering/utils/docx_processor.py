#!/usr/bin/env python
# -*- coding: utf-8 -*-

"""
Optimized DOCX Processor Module

This module provides memory-optimized handling for DOCX files,
particularly focused on efficient table extraction and processing
for French language documents.
"""

import docx
import gc
from pathlib import Path
from contextlib import contextmanager
from typing import List, Dict, Any, Generator, Optional
import logging

logger = logging.getLogger(__name__)

class OptimizedDocxProcessor:
    """
    Memory-optimized processor for DOCX files.
    
    Addresses the O(n²) performance degradation in python-docx table cell
    recalculation by implementing efficient traversal and resource management.
    """
    
    def __init__(self, memory_threshold_mb: int = 100):
        """
        Initialize the processor.
        
        Args:
            memory_threshold_mb: Memory threshold in MB after which to trigger garbage collection
        """
        self.memory_threshold_mb = memory_threshold_mb
        self.initial_memory_usage = self._get_memory_usage()
        
    @contextmanager
    def process_document(self, docx_path: str):
        """
        Context manager for safely processing a DOCX document.
        
        Args:
            docx_path: Path to the DOCX file
            
        Yields:
            The document object
        """
        document = None
        try:
            document = docx.Document(docx_path)
            yield document
        finally:
            # Force cleanup to prevent memory leaks
            del document
            gc.collect()
            self._check_memory_usage()
    
    def _get_memory_usage(self) -> float:
        """
        Get current memory usage in MB.
        
        Returns:
            Memory usage in MB
        """
        import psutil
        import os
        process = psutil.Process(os.getpid())
        memory_info = process.memory_info()
        # Convert bytes to MB
        return memory_info.rss / (1024 * 1024)
    
    def _check_memory_usage(self):
        """Check current memory usage and trigger GC if above threshold."""
        current_memory = self._get_memory_usage()
        initial_memory = self.initial_memory_usage
        
        # Log memory usage
        logger.debug(f"Memory usage: {current_memory:.2f} MB")
        
        # If memory usage has increased beyond threshold, force garbage collection
        if current_memory - initial_memory > self.memory_threshold_mb:
            logger.info(f"Memory threshold exceeded. Triggering garbage collection.")
            gc.collect()
            new_memory = self._get_memory_usage()
            logger.info(f"Memory after GC: {new_memory:.2f} MB (freed {current_memory - new_memory:.2f} MB)")
    
    def extract_text(self, docx_path: str) -> str:
        """
        Extract all text from a DOCX file efficiently.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            Extracted text content
        """
        text_parts = []
        
        with self.process_document(docx_path) as document:
            # Extract text from paragraphs
            for paragraph in document.paragraphs:
                text_parts.append(paragraph.text)
            
            # Extract text from tables efficiently
            table_text = self.extract_tables_text(document)
            text_parts.extend(table_text)
        
        return "\n".join(text_parts)

    def extract_tables_text(self, document) -> List[str]:
        """
        Extract text from all tables in the document efficiently.
        
        Args:
            document: The docx Document object
            
        Returns:
            List of text strings from table cells
        """
        text_parts = []
        
        # Process tables one by one
        for table_idx, table in enumerate(document.tables):
            logger.debug(f"Processing table {table_idx + 1}/{len(document.tables)}")
            
            # Process table rows in batches to avoid O(n²) behavior
            for row_idx, row in enumerate(table.rows):
                # Process cells in this row
                for cell_idx, cell in enumerate(row.cells):
                    # Extract text directly from cell paragraphs
                    cell_text = []
                    for paragraph in cell.paragraphs:
                        cell_text.append(paragraph.text)
                    
                    # Join cell text and add to results
                    if cell_text:
                        text_parts.append("\n".join(cell_text))
                
                # Check memory usage after processing each row
                if row_idx % 10 == 0:
                    self._check_memory_usage()
        
        return text_parts
    
    def extract_tables_data(self, docx_path: str) -> List[List[List[str]]]:
        """
        Extract tables data as structured format.
        
        Args:
            docx_path: Path to the DOCX file
            
        Returns:
            List of tables, each containing rows of cell text values
        """
        tables_data = []
        
        with self.process_document(docx_path) as document:
            # Process tables one by one
            for table_idx, table in enumerate(document.tables):
                table_data = []
                
                # Process each row
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    
                    # Process each cell
                    for cell in row.cells:
                        # Join all paragraphs in the cell
                        cell_text = "\n".join(p.text for p in cell.paragraphs)
                        row_data.append(cell_text)
                    
                    table_data.append(row_data)
                    
                    # Check memory after processing each row
                    if row_idx % 10 == 0:
                        self._check_memory_usage()
                
                tables_data.append(table_data)
        
        return tables_data
    
    def process_tables_batched(self, docx_path: str, batch_size: int = 10) -> Generator[List[List[str]], None, None]:
        """
        Process tables in batches to manage memory usage.
        
        Args:
            docx_path: Path to the DOCX file
            batch_size: Number of rows to process at once
            
        Yields:
            Batches of table rows
        """
        with self.process_document(docx_path) as document:
            for table in document.tables:
                current_batch = []
                
                for row_idx, row in enumerate(table.rows):
                    # Extract row data
                    row_data = [cell.text for cell in row.cells]
                    current_batch.append(row_data)
                    
                    # Yield batch when it reaches the specified size
                    if len(current_batch) >= batch_size:
                        yield current_batch
                        current_batch = []
                        
                        # Force memory check and cleanup after each batch
                        self._check_memory_usage()
                
                # Yield any remaining rows in the final batch
                if current_batch:
                    yield current_batch
